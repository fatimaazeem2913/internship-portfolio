import os
from typing import List
from pypdf import PdfReader
import docx
from langchain_core.documents import Document


class DocumentIngestionEngine:
    """Parses and chunks heterogeneous document formats (PDF, DOCX, TXT, MD)."""
    
    def __init__(self, chunk_size_words: int = 180, chunk_overlap_words: int = 35):
        self.chunk_size_words = chunk_size_words
        self.chunk_overlap_words = chunk_overlap_words

    def parse_pdf(self, file_path: str, filename: str) -> List[Document]:
        docs = []
        try:
            reader = PdfReader(file_path)
            for idx, page in enumerate(reader.pages):
                text = page.extract_text() or ""
                clean_text = text.strip()
                if not clean_text:
                    continue
                
                words = clean_text.split()
                if len(words) <= self.chunk_size_words:
                    docs.append(Document(
                        page_content=clean_text,
                        metadata={
                            "source": filename,
                            "page": idx + 1,
                            "chunk_id": f"{filename}_p{idx+1}_c0",
                            "doc_type": "pdf"
                        }
                    ))
                else:
                    step = self.chunk_size_words - self.chunk_overlap_words
                    for c_idx, start in enumerate(range(0, len(words), step)):
                        chunk_words = words[start:start + self.chunk_size_words]
                        if not chunk_words:
                            continue
                        chunk_text = " ".join(chunk_words)
                        docs.append(Document(
                            page_content=chunk_text,
                            metadata={
                                "source": filename,
                                "page": idx + 1,
                                "chunk_id": f"{filename}_p{idx+1}_c{c_idx}",
                                "doc_type": "pdf"
                            }
                        ))
        except Exception as e:
            print(f"Error parsing PDF {filename}: {e}")
        return docs

    def parse_docx(self, file_path: str, filename: str) -> List[Document]:
        docs = []
        try:
            doc = docx.Document(file_path)
            full_text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
            words = full_text.split()
            step = self.chunk_size_words - self.chunk_overlap_words
            for c_idx, start in enumerate(range(0, len(words), step)):
                chunk_words = words[start:start + self.chunk_size_words]
                if not chunk_words:
                    continue
                docs.append(Document(
                    page_content=" ".join(chunk_words),
                    metadata={
                        "source": filename,
                        "page": 1,
                        "chunk_id": f"{filename}_p1_c{c_idx}",
                        "doc_type": "docx"
                    }
                ))
        except Exception as e:
            print(f"Error parsing DOCX {filename}: {e}")
        return docs

    def parse_text(self, file_path: str, filename: str) -> List[Document]:
        docs = []
        try:
            with open(file_path, "r", encoding="utf-8") as f:
                content = f.read()
            words = content.split()
            step = self.chunk_size_words - self.chunk_overlap_words
            for c_idx, start in enumerate(range(0, len(words), step)):
                chunk_words = words[start:start + self.chunk_size_words]
                if not chunk_words:
                    continue
                docs.append(Document(
                    page_content=" ".join(chunk_words),
                    metadata={
                        "source": filename,
                        "page": 1,
                        "chunk_id": f"{filename}_p1_c{c_idx}",
                        "doc_type": "txt"
                    }
                ))
        except Exception as e:
            print(f"Error parsing text file {filename}: {e}")
        return docs

    def process_file(self, file_path: str, filename: str) -> List[Document]:
        ext = os.path.splitext(filename)[1].lower()
        if ext == ".pdf":
            return self.parse_pdf(file_path, filename)
        elif ext == ".docx":
            return self.parse_docx(file_path, filename)
        elif ext in [".txt", ".md", ".json"]:
            return self.parse_text(file_path, filename)
        return []