"""One-off script: generates a real DOCX file ('data/docx/product_spec.docx')
with real heading levels, paragraphs, and a table -- so python-docx has
genuine document structure to extract, not just flat text."""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

doc.add_heading("Acme Logistics Platform — Product Specification", level=0)
doc.add_paragraph("Document Version: 2.3 | Last Updated: January 2026")

doc.add_heading("1. Overview", level=1)
doc.add_paragraph(
    "The Acme Logistics Platform is a cloud-based system for route optimization, "
    "shipment tracking, and warehouse inventory management. This document specifies "
    "the core functional requirements for the platform's v2.3 release, targeted at "
    "mid-market logistics companies operating 5-50 vehicle fleets."
)

doc.add_heading("2. Route Optimization Module", level=1)
doc.add_heading("2.1 Optimization Algorithm", level=2)
doc.add_paragraph(
    "The route optimization engine uses a constrained vehicle routing algorithm that "
    "accounts for delivery time windows, vehicle capacity, and driver working-hour "
    "limits. Routes are recalculated automatically whenever a new order is added within "
    "30 minutes of the current dispatch window, and manually on-demand at any other time."
)
doc.add_heading("2.2 Real-Time Traffic Integration", level=2)
doc.add_paragraph(
    "Route calculations incorporate real-time traffic data refreshed every 5 minutes "
    "during active delivery windows. If traffic data becomes unavailable, the system "
    "falls back to historical average travel times for the given route segment and "
    "time of day, and flags affected routes as 'estimated' in the dispatcher UI."
)

doc.add_heading("3. Shipment Tracking Module", level=1)
doc.add_heading("3.1 Status Update Frequency", level=2)
doc.add_paragraph(
    "Shipment status updates are pushed to the customer-facing tracking portal within "
    "60 seconds of a driver marking a delivery event (picked up, in transit, delivered, "
    "or exception) in the driver mobile app. Customers receive an SMS notification for "
    "delivery and exception events only, not for routine in-transit updates."
)
doc.add_heading("3.2 Exception Handling", level=2)
doc.add_paragraph(
    "Delivery exceptions (failed delivery, damaged package, customer unavailable) "
    "require the driver to select a reason code and optionally attach a photo before "
    "the exception can be submitted. Exceptions automatically create a follow-up task "
    "assigned to the dispatcher on duty, due within 4 business hours."
)

doc.add_heading("4. Warehouse Inventory Module", level=1)
doc.add_heading("4.1 Stock Level Thresholds", level=2)
doc.add_paragraph(
    "Each SKU has a configurable low-stock threshold. When on-hand quantity falls below "
    "this threshold, the system generates a replenishment alert visible on the warehouse "
    "manager's dashboard and, if enabled, sends an automated email to the designated "
    "procurement contact for that SKU."
)
doc.add_heading("4.2 Cycle Counting", level=2)
doc.add_paragraph(
    "The platform supports scheduled cycle counts at the SKU or bin-location level. "
    "Discrepancies between system-recorded and physically counted quantities exceeding "
    "a 2% variance threshold automatically flag the SKU for a full recount before the "
    "adjusted quantity is accepted into the system of record."
)

doc.add_heading("5. Non-Functional Requirements", level=1)

table = doc.add_table(rows=1, cols=3)
table.style = "Light Grid Accent 1"
hdr = table.rows[0].cells
hdr[0].text = "Requirement"
hdr[1].text = "Target"
hdr[2].text = "Measurement Method"

rows = [
    ("API response time (p95)", "< 400ms", "Synthetic monitoring, 5-minute intervals"),
    ("Platform uptime (monthly)", "99.9%", "Third-party uptime monitoring service"),
    ("Route recalculation time", "< 90 seconds", "Internal application logging"),
    ("Mobile app crash rate", "< 0.5% of sessions", "Mobile crash reporting SDK"),
]
for req, target, method in rows:
    cells = table.add_row().cells
    cells[0].text = req
    cells[1].text = target
    cells[2].text = method

doc.add_heading("6. Out of Scope for v2.3", level=1)
doc.add_paragraph(
    "The following capabilities are explicitly out of scope for this release and are "
    "tracked separately on the v3.0 roadmap: multi-warehouse transfer optimization, "
    "predictive maintenance alerts for fleet vehicles, and customer-facing route ETA "
    "sharing via public API. Any customer requests for these capabilities should be "
    "logged in the product feedback tracker rather than treated as a defect."
)

doc.save("data/docx/product_spec.docx")
print("DOCX generated.")
