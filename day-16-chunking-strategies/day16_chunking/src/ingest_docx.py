"""
ingest_docx.py
--------------
DOCX ingestion using python-docx.

DOCX has no concept of "pages" the way PDF does (page breaks are a
rendering-time concept, not stored per-paragraph), so instead of a page
number, each extracted block carries:
  - its heading level (0 = body text, 1 = Heading 1, 2 = Heading 2, ...)
  - the nearest preceding heading text, used later as "section heading"
    metadata on chunks -- this is what lets a DOCX chunk cite a meaningful
    section reference despite having no page number.
Tables are extracted separately and flattened into row-per-line text so
their content isn't silently dropped.
"""

from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path

from docx import Document


@dataclass
class DocxBlock:
    source: str
    block_index: int          # order of appearance in the document
    text: str
    heading_level: int        # 0 = body paragraph, 1+ = heading of that level
    section_heading: str      # nearest preceding heading, used as metadata


def _heading_level(paragraph) -> int:
    style_name = paragraph.style.name if paragraph.style else ""
    if style_name.startswith("Heading"):
        try:
            return int(style_name.split(" ")[-1])
        except ValueError:
            return 1
    if style_name == "Title":
        return 0  # treated as the document title, not a numbered section
    return -1  # not a heading at all -- body text


def extract_docx(path: str) -> list[DocxBlock]:
    source = Path(path).name
    doc = Document(path)
    blocks: list[DocxBlock] = []
    current_heading = "Document Start"

    idx = 0
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        level = _heading_level(para)
        if level >= 1:
            current_heading = text
            blocks.append(DocxBlock(source, idx, text, heading_level=level, section_heading=text))
        else:
            blocks.append(DocxBlock(source, idx, text, heading_level=0, section_heading=current_heading))
        idx += 1

    # Tables: flatten each row into a single pipe-delimited line so the
    # information isn't lost, tagged with the section heading active at
    # the point the table appears in the document body.
    for table in doc.tables:
        table_lines = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            table_lines.append(" | ".join(cells))
        table_text = "\n".join(table_lines)
        blocks.append(DocxBlock(source, idx, table_text, heading_level=0, section_heading=f"{current_heading} (table)"))
        idx += 1

    return blocks


if __name__ == "__main__":
    blocks = extract_docx("data/docx/product_spec.docx")
    print(f"Extracted {len(blocks)} blocks from DOCX.\n")
    for b in blocks[:8]:
        tag = f"[H{b.heading_level}]" if b.heading_level >= 1 else "[body]"
        print(f"{tag} (section: {b.section_heading[:40]}) {b.text[:80]}")

    table_blocks = [b for b in blocks if "(table)" in b.section_heading]
    print(f"\nTable blocks found: {len(table_blocks)}")
    if table_blocks:
        print("Sample table block:\n" + table_blocks[0].text)
